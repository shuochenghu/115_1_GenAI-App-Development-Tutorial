"""第 10 週沿用的本機文件讀取、清理與 chunking helper。

這個模組延續第 9 週正式教材的前處理管線，只負責把不同格式檔案轉成可切塊的
文字。Embedding、語意搜尋與付費 API 呼叫放在 `embedding_utils.py` 與 `app.py`，
讓學生能清楚分辨「文件處理」和「向量搜尋」兩層責任。
"""

from __future__ import annotations

from io import BytesIO
import re

import pandas as pd
from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = ("pdf", "docx", "csv", "txt", "md")


def decode_text_bytes(data: bytes) -> str:
    """將純文字 bytes 解碼成 Python 字串。

    參數：
        data: 從上傳檔案取得的原始 bytes。

    回傳：
        解碼後的文字內容。

    可能錯誤：
        ValueError: 常見 UTF-8 與繁中 Windows 編碼都失敗時拋出。
    """
    for encoding in ("utf-8-sig", "utf-8", "cp950"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("無法辨識文字編碼，請將檔案另存為 UTF-8。")


def extract_pdf_text(data: bytes) -> str:
    """逐頁抽取 PDF 文字，並保留頁碼作為後續搜尋結果的來源線索。"""
    reader = PdfReader(BytesIO(data))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"[第 {page_number} 頁]\n{page_text.strip()}")
    if not pages:
        raise ValueError("PDF 沒有可抽取文字；它可能是掃描檔，需要 OCR。")
    return "\n\n".join(pages)


def extract_docx_text(data: bytes) -> str:
    """讀取 Word 文件中的段落與表格文字，轉成統一字串。"""
    document = Document(BytesIO(data))
    blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table_index, table in enumerate(document.tables, start=1):
        blocks.append(f"[表格 {table_index}]")
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    if not blocks:
        raise ValueError("Word 文件沒有可讀取的段落或表格文字。")
    return "\n".join(blocks)


def extract_csv_text(data: bytes, max_rows: int = 200) -> str:
    """將 CSV 前幾列轉成帶欄位名稱的文字，避免大型表格直接送進模型。"""
    text = decode_text_bytes(data)
    frame = pd.read_csv(BytesIO(text.encode("utf-8")))
    if frame.empty:
        raise ValueError("CSV 沒有資料列。")

    limited = frame.head(max_rows).fillna("")
    lines = [f"欄位：{', '.join(map(str, limited.columns))}"]
    for index, row in limited.iterrows():
        fields = [f"{column}={row[column]}" for column in limited.columns]
        lines.append(f"第 {index + 1} 列：" + "；".join(fields))
    if len(frame) > max_rows:
        lines.append(f"[僅載入前 {max_rows} 列；原檔共有 {len(frame)} 列]")
    return "\n".join(lines)


def extract_text(filename: str, data: bytes) -> str:
    """依副檔名選擇本機 reader，統一回傳後續管線可使用的文字。"""
    if "." not in filename:
        raise ValueError("檔名沒有副檔名。")
    suffix = filename.lower().rsplit(".", maxsplit=1)[-1]
    readers = {
        "txt": decode_text_bytes,
        "md": decode_text_bytes,
        "pdf": extract_pdf_text,
        "docx": extract_docx_text,
        "csv": extract_csv_text,
    }
    if suffix not in readers:
        raise ValueError(f"不支援 .{suffix}；請使用 PDF、DOCX、CSV、TXT 或 MD。")
    return readers[suffix](data)


def clean_text(text: str) -> str:
    """標準化空白與換行，同時保留段落邊界。"""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in normalized.split("\n")]
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 120) -> list[dict]:
    """用固定長度滑動視窗切分文字，並記錄每段來源位置。"""
    if chunk_size <= 0:
        raise ValueError("chunk_size 必須大於 0。")
    if overlap < 0 or overlap >= chunk_size:
        raise ValueError("overlap 必須介於 0（含）與 chunk_size（不含）之間。")

    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        content = text[start:end].strip()
        if content:
            chunks.append({
                "chunk_id": len(chunks),
                "start": start,
                "end": end,
                "text": content,
            })
        if end == len(text):
            break
        start = end - overlap
    return chunks
