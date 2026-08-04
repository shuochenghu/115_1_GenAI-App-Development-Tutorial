"""第 9 週本機文件讀取、清理與 chunking helper。

這個模組不依賴 Streamlit，也不呼叫 OpenAI API。課堂上可以先獨立測試這些
純資料處理函式，再接到 `app.py` 的 UI，讓學生分辨「本機前處理」與「AI 呼叫」。
"""

from __future__ import annotations

from io import BytesIO
import re

import pandas as pd
from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = ("pdf", "docx", "csv", "txt", "md")


def decode_text_bytes(data: bytes) -> str:
    """將純文字 bytes 解碼成 Python 字串。

    參數：
        data: 從上傳檔案取得的原始 bytes。

    回傳：
        解碼後的文字內容。

    可能錯誤：
        ValueError: 常見 UTF-8 與繁中 Windows 編碼都失敗時拋出，提醒學生轉存。

    教學重點：
    - `utf-8-sig` 可處理帶 BOM 的文字檔。
    - `cp950` 是繁體中文 Windows 常見編碼，放在 UTF-8 後面作為備援選項。
    """
    # 依序嘗試常見編碼；成功就立刻回傳，失敗才換下一個。
    for encoding in ("utf-8-sig", "utf-8", "cp950"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("無法辨識文字編碼，請將檔案另存為 UTF-8。")


def extract_pdf_text(data: bytes) -> str:
    """逐頁抽取 PDF 文字，並保留頁碼作為來源線索。

    參數：
        data: PDF 檔案的原始 bytes。

    回傳：
        合併後的文字；每頁前面會加上 `[第 N 頁]`，方便後續摘要回查來源。

    可能錯誤：
        ValueError: PDF 沒有文字層時拋出，常見原因是掃描圖檔 PDF，需要 OCR。
    """
    reader = PdfReader(BytesIO(data))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            # 保留頁碼比單純串接文字更適合教學，學生能追蹤 chunk 來自哪一頁。
            pages.append(f"[第 {page_number} 頁]\n{page_text.strip()}")
    if not pages:
        raise ValueError("PDF 沒有可抽取文字；它可能是掃描檔，需要 OCR。")
    return "\n\n".join(pages)


def extract_docx_text(data: bytes) -> str:
    """讀取 Word 文件中的段落與表格文字。

    參數：
        data: DOCX 檔案的原始 bytes。

    回傳：
        依文件順序整理出的文字；表格會以 `|` 分隔儲存格，讓內容仍可閱讀。

    可能錯誤：
        ValueError: 文件沒有可讀取的段落或表格文字時拋出。
    """
    document = Document(BytesIO(data))
    # 先處理一般段落；空段落通常只是排版，不適合進入摘要或 chunking。
    blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table_index, table in enumerate(document.tables, start=1):
        # 表格先加標記，避免學生在預覽時分不清楚段落文字與表格資料。
        blocks.append(f"[表格 {table_index}]")
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    if not blocks:
        raise ValueError("Word 文件沒有可讀取的段落或表格文字。")
    return "\n".join(blocks)


def extract_csv_text(data: bytes, max_rows: int = 200) -> str:
    """將 CSV 前幾列轉成帶欄位名稱的文字。

    參數：
        data: CSV 檔案的原始 bytes。
        max_rows: 最多轉換的列數；避免大型表格讓 notebook 或 app 變慢。

    回傳：
        可直接預覽、切塊或摘要的文字表示。

    可能錯誤：
        ValueError: CSV 沒有任何資料列時拋出。
    """
    text = decode_text_bytes(data)
    frame = pd.read_csv(BytesIO(text.encode("utf-8")))
    if frame.empty:
        raise ValueError("CSV 沒有資料列。")
    # 課堂先用前 max_rows 列建立可控範例；完整表格分析留到後續資料處理課題。
    limited = frame.head(max_rows).fillna("")
    lines = [f"欄位：{', '.join(map(str, limited.columns))}"]
    for index, row in limited.iterrows():
        # 每列保留欄名，讓 chunk 被單獨拿出來時仍看得懂欄位語意。
        fields = [f"{column}={row[column]}" for column in limited.columns]
        lines.append(f"第 {index + 1} 列：" + "；".join(fields))
    if len(frame) > max_rows:
        lines.append(f"[僅載入前 {max_rows} 列；原檔共有 {len(frame)} 列]")
    return "\n".join(lines)


def extract_text(filename: str, data: bytes) -> str:
    """依副檔名選擇合適的本機 reader，統一回傳文字。

    參數：
        filename: 使用者上傳檔案的原始檔名，用來判斷副檔名。
        data: 上傳檔案內容的原始 bytes。

    回傳：
        各格式 reader 抽取後的文字內容。

    可能錯誤：
        ValueError: 檔名沒有副檔名，或副檔名不在本週支援範圍。

    教學重點：
    - UI 不應直接寫一大串 if/elif；把格式路由集中在這個函式更容易測試。
    - 所有讀取器都回傳 `str`，後面的清理與 chunking 才能共用同一條管線。
    """
    if "." not in filename:
        raise ValueError("檔名沒有副檔名。")
    suffix = filename.lower().rsplit(".", maxsplit=1)[-1]
    # 字典把「副檔名」映射到「處理函式」，新增格式時只需擴充這裡。
    readers = {
        "txt": decode_text_bytes,
        "md": decode_text_bytes,
        "pdf": extract_pdf_text,
        "docx": extract_docx_text,
        "csv": extract_csv_text,
    }
    if suffix not in readers:
        raise ValueError(f"不支援 .{suffix}；請使用 PDF、DOCX、CSV、TXT 或 MD。")
    return readers[suffix](data)


def clean_text(text: str) -> str:
    """標準化空白與換行，同時保留段落邊界。

    參數：
        text: reader 抽取出的原始文字，可能混有 Windows/macOS 換行與多餘空白。

    回傳：
        清理後的文字；行內多空白會合併，三個以上連續換行會壓成段落分隔。

    教學重點：
    - 清理不是把所有換行刪掉；段落邊界常是摘要與 chunking 的重要語意線索。
    - 先把不同作業系統的換行統一成 `\n`，後續規則才會穩定。
    """
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    # 每一行只清行內空白，不急著移除段落換行。
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in normalized.split("\n")]
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 120) -> list[dict]:
    """用固定長度滑動視窗切分文字，並記錄每段來源位置。

    參數：
        text: 已清理完成的文件文字。
        chunk_size: 每個 chunk 最多保留的字元數。
        overlap: 相鄰 chunk 重疊的字元數，用來保留交界處上下文。

    回傳：
        `list[dict]`，每個 dict 包含 `chunk_id`、`start`、`end` 與 `text`。

    可能錯誤：
        ValueError: chunk_size 或 overlap 不合理，可能造成空 chunk 或無限迴圈。

    教學重點：
    - 這是第 10 週 embedding / RAG 的前置模型：先能切出可追蹤的 chunks。
    - `start` / `end` 是原文位置，之後可用來回查來源或除錯切塊結果。
    """
    if chunk_size <= 0:
        raise ValueError("chunk_size 必須大於 0。")
    if overlap < 0 or overlap >= chunk_size:
        raise ValueError("overlap 必須介於 0（含）與 chunk_size（不含）之間。")

    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        content = text[start:end].strip()
        if content:
            # 保留字元範圍是教學與除錯用中繼資料，不只是顯示文字。
            chunks.append({
                "chunk_id": len(chunks),
                "start": start,
                "end": end,
                "text": content,
            })
        if end == len(text):
            break
        # 往回 overlap 個字元，讓下一段保留上一段結尾的語境。
        start = end - overlap
    return chunks
