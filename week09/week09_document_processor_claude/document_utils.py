"""第 9 週｜本機文件讀取、清理與 chunking 工具（Claude 版）。

這個模組只做「送進模型之前」的確定性資料處理，完全不呼叫任何付費 API：

    上傳位元組 → 依副檔名路由 reader → 抽取文字 → 清理 → 切成 chunks

app.py 只有在使用者按下按鈕時，才會拿清理後的文字去呼叫 OpenAI。
"""

from __future__ import annotations

from io import BytesIO
import re

import pandas as pd
from docx import Document
from pypdf import PdfReader


# App 允許上傳的副檔名（小寫、不含點）。
SUPPORTED_EXTENSIONS = ("pdf", "docx", "csv", "txt", "md")


def decode_text_bytes(data: bytes) -> str:
    """依常見編碼順序解碼純文字；全部失敗時給出明確錯誤訊息。"""
    for encoding in ("utf-8-sig", "utf-8", "cp950"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("無法辨識文字編碼，請將檔案另存為 UTF-8 後再試。")


def extract_pdf_text(data: bytes) -> str:
    """逐頁抽取 PDF 文字，保留頁碼標記以利之後追蹤來源。"""
    reader = PdfReader(BytesIO(data))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"[第 {page_number} 頁]\n{page_text.strip()}")
    if not pages:
        raise ValueError("PDF 沒有可抽取的文字，它可能是掃描檔，需要 OCR。")
    return "\n\n".join(pages)


def extract_docx_text(data: bytes) -> str:
    """讀取 Word 段落與表格，轉成統一文字。"""
    document = Document(BytesIO(data))
    blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table_index, table in enumerate(document.tables, start=1):
        blocks.append(f"[表格 {table_index}]")
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    if not blocks:
        raise ValueError("Word 文件沒有可讀取的段落或表格文字。")
    return "\n".join(blocks)


def extract_csv_text(data: bytes, max_rows: int = 200) -> str:
    """保留欄名，將 CSV 前 max_rows 列轉成適合摘要的文字。"""
    text = decode_text_bytes(data)
    frame = pd.read_csv(BytesIO(text.encode("utf-8")))
    if frame.empty:
        raise ValueError("CSV 沒有資料列。")
    limited = frame.head(max_rows).fillna("")
    lines = [f"欄位：{', '.join(map(str, limited.columns))}"]
    for index, row in limited.iterrows():
        fields = [f"{column}={row[column]}" for column in limited.columns]
        lines.append(f"第 {index + 1} 列：" + "；".join(fields))
    if len(frame) > max_rows:
        lines.append(f"[僅載入前 {max_rows} 列；原檔共有 {len(frame)} 列]")
    return "\n".join(lines)


def extract_text(filename: str, data: bytes) -> str:
    """依副檔名把上傳檔案路由到正確的 reader，回傳統一文字字串。"""
    if "." not in filename:
        raise ValueError("檔名沒有副檔名，無法判斷格式。")
    suffix = filename.lower().rsplit(".", maxsplit=1)[-1]
    readers = {
        "txt": decode_text_bytes,
        "md": decode_text_bytes,
        "pdf": extract_pdf_text,
        "docx": extract_docx_text,
        "csv": extract_csv_text,
    }
    if suffix not in readers:
        raise ValueError(f"不支援 .{suffix}；請使用 PDF、DOCX、CSV、TXT 或 MD。")
    return readers[suffix](data)


def clean_text(text: str) -> str:
    """統一換行、移除行內多餘空白，並保留段落分隔（連續空行縮成一個空行）。"""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in normalized.split("\n")]
    cleaned = "\n".join(lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 120) -> list[dict]:
    """用固定長度的滑動視窗切分文字。

    每個 chunk 是一個 dict，含 chunk_id、start、end 與 text，
    方便第 10、11 週追蹤來源位置。相鄰兩塊會重疊 overlap 個字元，
    降低句子剛好被切斷造成語意遺失的風險。
    """
    if chunk_size <= 0:
        raise ValueError("chunk_size 必須大於 0。")
    if overlap < 0 or overlap >= chunk_size:
        raise ValueError("overlap 必須介於 0（含）與 chunk_size（不含）之間。")

    chunks: list[dict] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        content = text[start:end].strip()
        if content:
            chunks.append({
                "chunk_id": len(chunks),
                "start": start,
                "end": end,
                "text": content,
            })
        if end == len(text):
            break
        start = end - overlap
    return chunks
